"""Generate .docx plan for Face Scan + Valid ID Verification feature."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "face_scan_verification_plan.docx")

PRIMARY = RGBColor(0x0E, 0xA5, 0xE9)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x64, 0x74, 0x8B)
GREEN = RGBColor(0x16, 0x65, 0x34)
RED = RGBColor(0xB9, 0x1C, 0x1C)

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def heading(text, level=1, color=DARK):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"


def para(text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)


def bullet(text, prefix=None, color=None):
    p = doc.add_paragraph(style="List Bullet")
    if prefix:
        r = p.add_run(prefix)
        r.bold = True
        if color:
            r.font.color.rgb = color
    r = p.add_run(text)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)


def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F5F9")
    pPr.append(shd)


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "0EA5E9")
        tcPr.append(shd)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.rows[r + 1].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            if c == 0:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
    doc.add_paragraph()


# ===== TITLE PAGE =====
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PureDrop Capstone Project")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Feature Plan: Face Scan + Valid ID Verification")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Post-Registration Identity Verification Flow\nwith Real-Time Liveness Detection")
r.font.size = Pt(14)
r.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Date: August 11, 2026\nStatus: Proposed / For Implementation\nPlatform: React Native (Expo) + Firebase + Supabase")
r.font.size = Pt(11)
r.font.color.rgb = GRAY
doc.add_page_break()

# ===== 1. OVERVIEW =====
heading("1. Overview", 1, PRIMARY)
para("This document outlines the plan to add a two-step identity verification process after a user successfully registers and verifies their email. The goal is to ensure that every registered account belongs to a REAL person by requiring:")
bullet("A live selfie verified through real-time liveness detection (eye-open, head-turn, smile) plus a server-side Sightengine check (not AI-generated, not a screenshot, not a deepfake).", "Step 1 — Face Scan with Liveness: ")
bullet("A photo of a valid government-issued ID for manual or automated review.", "Step 2 — Valid ID Upload: ")
para("This feature is designed to be implemented at zero additional cost by combining free open-source on-device face detection with the existing Sightengine integration (already used for report attachment authenticity) and the existing Supabase Storage infrastructure.")

# ===== 2. FLOW =====
heading("2. Registration & Verification Flow", 1, PRIMARY)
code("Register → Email OTP → ✅ Email Verified\n    ↓\n🔹 Step 1: Face Scan (LIVE + REAL-TIME)\n   ├── Camera opens with face oval frame guide\n   ├── \"Look at the camera\" → eyes open detected\n   ├── \"Turn head left\" → head yaw detected\n   ├── \"Turn head right\" → head yaw detected\n   ├── \"Smile to verify\" → smile detected\n   ├── ✅ Liveness PASSED → auto-capture selfie\n   └── → Sightengine server check (deepfake/genai/type)\n    ↓\n🔹 Step 2: Valid ID Upload (Front + Back)\n    ↓\n🎉 Verification Complete → Login")

# ===== 3. TWO-LAYER =====
heading("3. Two-Layer Anti-Spoofing Architecture", 1, PRIMARY)

heading("3.1 Layer 1 — On-Device Real-Time Liveness (FREE, offline)", 2)
para("Uses react-native-vision-camera-face-detector to analyze the live camera feed BEFORE allowing capture. This is the critical anti-spoofing layer:", )
table(
    ["Liveness Check", "What It Detects", "Defeats"],
    [
        ["Face detection in frame", "Real face bounding box + 106-point contours in the live feed", "Blank photos, objects, non-face images"],
        ["Eyes open probability", "Both eyes must be open above a threshold (leftEyeOpenProbability / rightEyeOpenProbability)", "Photos of closed eyes, sleeping people"],
        ["Head pose movement", "User is instructed to turn head left → right → center while yaw/pitch/roll angles are validated to change", "Printed photos, tablets, mannequins — static 2D images cannot produce real 3D head rotation"],
        ["Face contour tracking", "106-point contour map must track across frames", "Dolls, mannequins, wax statues — lack real facial geometry movement"],
        ["Smile detection", "smilingProbability must exceed a threshold", "Additional static-image proof"],
    ],
)

heading("3.2 Layer 2 — Server-Side Sightengine Check", 2)
para("After a successful live scan, the captured selfie is sent to a Supabase Edge Function that calls Sightengine:")
table(
    ["Model", "Purpose"],
    [
        ["faces", "Confirm a face is present in the captured photo"],
        ["deepfake", "Reject face-manipulated / deepfake images"],
        ["genai", "Reject AI-generated faces"],
        ["type", "Confirm it is a real photograph, not an illustration or screenshot"],
    ],
)

heading("3.3 Attack Scenario Defense", 2)
para("Without liveness detection, someone could hold up a printed photo, a tablet, a doll, or a mannequin and pass a static image check. Here is how the two layers stop each attack:")
table(
    ["Attack", "Layer 1: Liveness", "Layer 2: Sightengine"],
    [
        ["Printed photo of a person", "FAIL — no blinking, no head rotation", "May pass (face present)"],
        ["Tablet/laptop screen showing a face", "FAIL — static image, no 3D head movement", "May pass (face present)"],
        ["Mannequin / doll / wax figure", "FAIL — no eye openness, no facial muscle movement, contours don't track", "May pass (face present)"],
        ["AI-generated deepfake video", "MAY PASS (if using real-time video)", "FAIL — deepfake model detects manipulation"],
        ["AI-generated still image", "FAIL — no liveness responses", "FAIL — genai model detects AI generation"],
        ["Regular photo of a real person", "PASS — eyes open, head turns, smile", "PASS — real photograph"],
    ],
)

# ===== 4. DEPENDENCIES =====
heading("4. New Dependencies", 1, PRIMARY)
para("These open-source packages add native camera + face detection. They require a development build via EAS (already used in this project via expo-dev-client):")
code("npm install react-native-vision-camera\nnpm install react-native-vision-camera-face-detector @shopify/react-native-skia")
table(
    ["Package", "Purpose"],
    [
        ["react-native-vision-camera", "High-performance camera feed (replaces expo-image-picker camera for the face scan step)"],
        ["react-native-vision-camera-face-detector", "ML Kit / Vision Framework face tracking — free, on-device, works offline"],
        ["@shopify/react-native-skia", "Drawing the face contour overlay on the live preview"],
    ],
)
para("These run entirely ON-DEVICE for free — no API costs, no server round-trip, no Sightengine quota usage for the liveness check. Sightengine is only called AFTER a successful liveness pass, conserving the 100 free credits/month.", color=GRAY)

# ===== 5. STEP 2 =====
heading("5. Step 2: Valid ID Upload", 1, PRIMARY)
heading("5.1 How It Works", 2)
bullet("New screen: /login/verification/valid_id")
bullet("User takes photos of their valid government ID (front + back)")
bullet("Photos uploaded to Supabase Storage at:")
code("verification/{userId}/valid-id-front.jpg\nverification/{userId}/valid-id-back.jpg\nverification/{userId}/selfie.jpg")

heading("5.2 Firestore Document Update", 2)
para("The regular_user/{uid} document is updated with:")
code("{\n  verificationStatus: \"pending\",   // pending → verified / rejected\n  faceScanUrl: \"...\",               // selfie public URL\n  faceScanPath: \"...\",\n  livenessPassed: true,\n  validIdFrontUrl: \"...\",\n  validIdFrontPath: \"...\",\n  validIdBackUrl: \"...\",\n  validIdBackPath: \"...\",\n  verifiedAt: null\n}")

# ===== 6. ADMIN =====
heading("6. Admin Panel Changes", 1, PRIMARY)
bullet("Show user verification status badge (Pending / Verified / Rejected)")
bullet("Preview the selfie + ID photos for manual approval")
bullet("Approve / Reject buttons")
bullet("Filter users by verification status")
bullet("Show livenessPassed indicator for each user")

# ===== 7. FILES =====
heading("7. Files to Create / Modify", 1, PRIMARY)
table(
    ["File", "Action", "Purpose"],
    [
        ["components/verification/FaceScanCamera.tsx", "New", "VisionCamera live preview + face detector + oval guide + liveness state machine"],
        ["components/verification/LivenessOverlay.tsx", "New", "Skia drawing of face landmarks + progress prompts"],
        ["lib/verification/livenessDetector.ts", "New", "Logic for eye-open, head-turn, smile checks with thresholds"],
        ["supabase/functions/verify-user-face/index.ts", "New", "Sightengine check (faces, deepfake, genai, type) after liveness passes"],
        ["api/faceVerification.ts", "New", "Client API to call the edge function"],
        ["app/login/verification/face_scan.tsx", "New", "Face scan screen wrapper"],
        ["app/login/verification/valid_id.tsx", "New", "Valid ID upload screen"],
        ["app/login/verification/success.tsx", "New", "Verification complete screen"],
        ["lib/login/verificationStore.ts", "New", "Temporary store for captured face/ID URIs during flow"],
        ["lib/login/registerfunctions.ts", "Edit", "Add verificationStatus & verification fields to Firestore doc"],
        ["app/login/email_verification/success.tsx", "Edit", "Route to face scan instead of Login"],
        ["app/login/_layout.tsx", "Edit", "Register new screens in stack"],
        ["supabase/storage_policies_verification.sql", "New", "SQL policies for verification folder in regular_user bucket"],
        ["PureDop_Admin/src/admin/users/...", "Edit", "Show verification status + preview images + Approve/Reject"],
    ],
)

# ===== 8. QUESTIONS =====
heading("8. Open Questions", 1, PRIMARY)
bullet("Should the ID be reviewed by the admin (Pending → Approve/Reject), or just stored without a review step?")
bullet("Do you want to restrict to specific Philippine IDs (UMID, Driver's License, Passport, etc.) or allow any valid ID photo?")
bullet("Do you want to capture both front AND back of the ID, or just the front?")
bullet("Which liveness gestures are required: head-turn only, or head-turn + smile?")

# ===== 9. SUMMARY =====
heading("9. Summary", 1, PRIMARY)
para("This feature is achievable at zero cost by combining free on-device liveness detection (react-native-vision-camera-face-detector) with the existing Sightengine free tier + expo-image-picker + Supabase storage. The two-layer approach ensures that printed photos, tablets, mannequins, and AI-generated faces are all rejected — only a real, live human face passes.")

os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")